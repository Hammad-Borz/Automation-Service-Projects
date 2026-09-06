"""Text extraction for supported document formats."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from .exceptions import DocumentReaderError, EmptyDocumentError, UnsupportedDocumentError
from .models import DocumentMetadata, ExtractedDocument, PageText, new_id


class DocumentReader:
    """Extract text and metadata from PDF and DOCX files."""

    supported_extensions = {".pdf", ".docx"}

    def read(self, document_path: str | Path, document_id: str | None = None) -> ExtractedDocument:
        path = Path(document_path)
        if not path.is_file():
            raise FileNotFoundError(f"Document not found: {path}")
        suffix = path.suffix.lower()
        if suffix not in self.supported_extensions:
            raise UnsupportedDocumentError(
                f"Unsupported document type '{path.suffix}'. Supported types: PDF, DOCX."
            )

        try:
            pages = self._read_pdf(path) if suffix == ".pdf" else self._read_docx(path)
        except (UnsupportedDocumentError, EmptyDocumentError, DocumentReaderError):
            raise
        except (PdfReadError, OSError) as exc:
            raise DocumentReaderError(f"Could not read document '{path.name}'.") from exc
        except Exception as exc:
            raise DocumentReaderError(f"Could not read document '{path.name}'.") from exc

        combined = "\n".join(page.text for page in pages).strip()
        if not combined:
            raise EmptyDocumentError(f"Document '{path.name}' contains no extractable text.")

        metadata = DocumentMetadata(
            document_id=document_id or new_id("doc"),
            file_name=path.name,
            file_type=suffix.lstrip("."),
            source_path=str(path.resolve()),
            page_count=len(pages) or None,
        )
        return ExtractedDocument(metadata=metadata, text=combined, pages=pages)

    def extract_text(self, document_path: str | Path) -> str:
        """Compatibility helper used by tests that only need the text."""
        return self.read(document_path).text

    @staticmethod
    def _read_pdf(path: Path) -> list[PageText]:
        reader = PdfReader(str(path))
        pages: list[PageText] = []
        for index, page in enumerate(reader.pages, start=1):
            pages.append(PageText(page_number=index, text=(page.extract_text() or "").strip()))
        return pages

    @staticmethod
    def _read_docx(path: Path) -> list[PageText]:
        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        tables: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    tables.append(" | ".join(cells))
        text = "\n".join(part for part in [*paragraphs, *tables] if part and part.strip())
        return [PageText(page_number=None, text=text)]
