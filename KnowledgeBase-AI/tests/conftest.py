from pathlib import Path

import pytest
from docx import Document

from src.chunker import Chunker
from src.document_processor import DocumentProcessor
from src.document_reader import DocumentReader
from src.embeddings import HashingEmbeddingProvider
from src.knowledge_base import KnowledgeBase
from src.main import write_simple_pdf
from src.models import DocumentMetadata, ExtractedDocument
from src.vector_store import JsonVectorStore


@pytest.fixture
def tmp_kb(tmp_path: Path) -> KnowledgeBase:
    return KnowledgeBase(
        reader=DocumentReader(),
        processor=DocumentProcessor(),
        chunker=Chunker(chunk_size=120, chunk_overlap=20),
        embeddings=HashingEmbeddingProvider(dimension=64),
        store=JsonVectorStore(tmp_path / "index.json"),
    )


def write_docx(path: Path, text: str) -> Path:
    document = Document()
    for line in text.splitlines() or [text]:
        document.add_paragraph(line)
    document.save(path)
    return path


def sample_extracted(text: str, name: str = "policy.pdf") -> ExtractedDocument:
    return ExtractedDocument(
        metadata=DocumentMetadata(
            document_id="doc_test",
            file_name=name,
            file_type="pdf",
            source_path=name,
        ),
        text=text,
    )
